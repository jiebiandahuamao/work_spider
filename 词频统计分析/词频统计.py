#-*- coding:utf-8 -*-
import jieba
import docx
import datetime
import os #用来删除临时生成的docx
from win32com import client as wc  #需要先安装pypewin32库 和 word 软件
import sys
reload(sys)
sys.setdefaultencoding( "utf-8" )
#address =input("请输入小说的存放位置：")
 
 
def doc2Docx(filePath,fileName,fileType):
    #调用word软件，将doc转换为docx文件。
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(filePath + "\\" + fileName + "." + fileType)  # 目标路径下的文件
    doc.SaveAs(filePath + "\\" + fileName + ".docx", 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
    doc.Close()
    word.Quit()
 
def deleteFile(file):
    #删除文件
    if os.path.exists(file):
        os.remove(file)
    else:
        print ('不存在要删除的文件:%s' % file)
 
def readDocument(strFile):
    '''
    获取文档对象，将文档内容按段落读入，并存入doc中
    '''
    file = docx.Document(strFile)
    doc = ""
    for para in file.paragraphs:
        doc = doc + para.text
    # print (doc)
    return doc
 
def readTxt(strFile):
    try:
        file = open(strFile).read()
    except:
        file = open(strFile, encoding="utf-8").read()
    txt=file
    return txt
 
def segment(doc):
    '''
    用jieba分词对输入文档进行分词，并保存至本地（根据情况可跳过）
    '''
    seg_list = " ".join(jieba.cut(doc, cut_all=False)) #seg_list为str类型
    i=0
    # document_after_segment = open('分词结果.txt', 'a+')
    # document_after_segment.write(seg_list)
    # document_after_segment.close()
 
    return seg_list
 
def removeStopWords(seg_list):
    '''
    自行下载stopwords1893.txt停用词表，该函数实现去停用词
    '''
    wordlist_stopwords_removed = []
 
    stop_words = open('stopwords1893.txt')
    stop_words_text = stop_words.read()
 
    stop_words.close()
 
    stop_words_text_list = stop_words_text.split('\n')
    for i in range(len(stop_words_text_list)):
        stop_words_text_list[i]=stop_words_text_list[i].strip()
 
    # print (len(stop_words_text_list))
    # print (stop_words_text_list)
    after_seg_text_list = seg_list.split(' ')
    print (len(after_seg_text_list))
    for word in after_seg_text_list:
        if (word not in stop_words_text_list  and len(word)>=2):
            wordlist_stopwords_removed.append(word)
 
    print (len(wordlist_stopwords_removed))
    # without_stopwords = open('分词结果(去停用词).txt', 'a+')
    # without_stopwords.write(' '.join(wordlist_stopwords_removed))
    # without_stopwords.close()
    return ' '.join(wordlist_stopwords_removed)
 
def wordCount(segment_list,resultFileName):
    '''
        该函数实现词频的统计，并将统计结果存储至本地。
        在制作词云的过程中用不到，主要是在画词频统计图时用到。
    '''
    word_lst = []
    word_dict = {}
    with open(resultFileName+'词频统计(去停用词).txt','a+') as wf2:
        word_lst.append(segment_list.split(' '))
        for item in word_lst:
            for item2 in item:
                if item2 not in word_dict:
                    word_dict[item2] = 1
                else:
                    word_dict[item2] += 1
 
        word_dict_sorted = dict(sorted(word_dict.items(),key = lambda item:item[1], reverse=True))#按照词频从大到小排序
        for key in word_dict_sorted:
            wf2.write(key+' '+str(word_dict_sorted[key])+'\n')
    wf2.close()
 
    result=[]
    i=0
    keyOfTop=15
    for key in word_dict_sorted:
        result.append(key)
        i=i+1
        if(i>=keyOfTop):
            break
 
    return result
 
def collect(file,result):
    #分词结果汇总.txt是记录所有用该程序进行过分词操作的文件，里面存储着文件名和文件所对应的topK的分词。
    #该方法将一个文件的topK分词写入文件。
    fw=open("fz.txt","a")
    fw.write(file+",")
    for i in range(len(result)-1):
        fw.write(result[i]+",")
    fw.write(result[-1]+"\n")
    fw.close()
 
 
def main():
    
    filePath="C:/Users/28674/Desktop"
    fileName="aa"
    fileType="txt"
    filePath=filePath.replace(" ","")
    fileName=fileName.replace(" ","")
    fileType=fileType.replace(" ","")
    
    print filePath,fileName,fileType

    if(fileType=="doc"):
        doc2Docx(filePath,fileName,fileType)
        file=filePath + "\\" + fileName + ".docx"
        doc = readDocument(file)
        deleteFile(file)
    elif(fileType=="docx"):
        doc = readDocument(filePath + "\\" + fileName + "." + fileType)
    elif(fileType=="txt"):
        doc = readTxt(filePath + "\\" + fileName + "." + fileType)
 
    seg_list = segment(doc)
    seg_list2 = removeStopWords(seg_list)
    result= wordCount(seg_list2, filePath + "\\" + fileName)
    collect(fileName + "." + fileType,result)
main()
